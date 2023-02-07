import urllib.request
import datetime
import dateutil.relativedelta
import io
import pandas as pd
import numpy as np
import docx
import os
from utils import get_state_names, zip_output_files


class StateData(object):
    def __init__(self, state_name):
        self.state_name = state_name

    @staticmethod
    def _get_last_month() -> str:
        today = datetime.date.today()
        last_month = today - dateutil.relativedelta.relativedelta(months=1)
        last_month_str = last_month.strftime("%Y-%m-%d")

        return last_month_str

    def _get_last_months_url(self) -> str:

        last_month = self._get_last_month()
        clean_state_name = self.state_name.replace(' ', '%20')
        state_data_last_month_url = f'https://healthdata.gov/resource/j8mb-icvb.csv?$query=SELECT%0A%20%20%60state' \
                                    f'%60%2C%0A%20%20%60state_name%60%2C%0A%20%20%60state_fips%60%2C%0A%20%20%60fe' \
                                    f'ma_region%60%2C%0A%20%20%60overall_outcome%60%2C%0A%20%20%60date%60%2C%0A%20' \
                                    f'%20%60new_results_reported%60%2C%0A%20%20%60total_results_reported%60%2C%0A%' \
                                    f'20%20%60geocoded_state%60%0AWHERE%0A%20%20(%60state_name%60%20IN%20(%22' \
                                    f'{clean_state_name}%22))%0A%20%20AND%20(%60date%60%20%3E%20%22{last_month}T11%' \
                                    f'3A09%3A54%22%20%3A%3A%20floating_timestamp)%0AORDER%20BY%20%60date%60%20DESC' \
                                    f'%20NULL%20LAST%2C%20%60overall_outcome%60%20ASC%20NULL%20LAST'

        return state_data_last_month_url

    def get_last_months_data(self):
        print(self._get_last_months_url())
        data = urllib.request.urlopen(self._get_last_months_url())
        output = io.StringIO()
        for line in data:
            output.write(line.decode('ascii'))

        output.seek(0)
        output_df = pd.read_csv(output)

        return output_df

    def get_last_five_days_data(self):
        last_month_df = self.get_last_months_data()
        pos_neg_df = last_month_df[last_month_df['overall_outcome'].isin(['Positive', 'Negative'])]
        pos_neg_df['date'] = pd.to_datetime(pos_neg_df['date'])
        max_date = pos_neg_df['date'].max()
        cutoff_date = pos_neg_df["date"].max() - pd.Timedelta(days=5)
        mask = (pos_neg_df['date'] > cutoff_date) & (pos_neg_df['date'] <= max_date)

        return pos_neg_df[mask]

    def get_sorted_latest_5_days_data(self):
        five_days_df = self.get_last_five_days_data()
        table = pd.pivot_table(five_days_df, values=['new_results_reported', 'total_results_reported'],
                               index=['date'], columns=['overall_outcome'], aggfunc=np.sum)

        table.columns = [s1 + '_' + str(s2) for (s1, s2) in table.columns.tolist()]
        table['Date'] = table.index
        table = table[["Date", "new_results_reported_Positive", "new_results_reported_Negative",
                       "total_results_reported_Positive", "total_results_reported_Negative"]]
        table = table.rename(columns={"new_results_reported_Positive": "New Positive Tests Reported",
                              "new_results_reported_Negative": "New Negative Tests Reported",
                              "total_results_reported_Positive": "Total Positive Tests Reported",
                              "total_results_reported_Negative": "Total Negative Tests Reported"})

        return table


class WriteStateDataToWord(object):

    def __init__(self, state_name, output_folder):
        self.state_name = state_name
        self.state_data = StateData(self.state_name)
        self.output_folder = output_folder

    def create_output_folder(self):
        if not os.path.exists(self.output_folder):
            os.mkdir(self.output_folder)

    def write_data_to_word(self):
        self.create_output_folder()
        word_file_name = f'{self.output_folder}/{self.state_name}.docx'
        doc = docx.Document()
        today = str(datetime.date.today())
        heading = self.state_name + ', ' + today
        doc.add_heading(heading, 1)

        df = self.state_data.get_sorted_latest_5_days_data()
        t = doc.add_table(df.shape[0] + 1, df.shape[1])

        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]

        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i + 1, j).text = str(df.values[i, j])

        doc.save(word_file_name)


if __name__ == '__main__':

    state_names = get_state_names()
    output_folder = './output'
    for name in state_names:
        try:
            write_data = WriteStateDataToWord(name, output_folder)
            write_data.write_data_to_word()
        except Exception as e:
            print(f'Failed for state: {name}, Error:{e}')

    zip_output_files(output_folder)
